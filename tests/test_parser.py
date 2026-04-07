import io
import pytest
from unittest.mock import MagicMock
from app.parser import extract_text, extract_text_from_pdf, extract_text_from_docx


# ── Helpers ───────────────────────────────────────────────────
def make_uploaded_file(name: str, content: bytes) -> MagicMock:
    """
    Creates a fake Streamlit UploadedFile object.
    MagicMock lets us create an object with any attribute we want
    without needing a real Streamlit session running.
    """
    mock = MagicMock()
    mock.name = name
    mock.read.return_value = content
    return mock


# ── Tests ─────────────────────────────────────────────────────
def test_extract_text_unsupported_format():
    """Uploading a .txt file should raise a ValueError."""
    mock_file = make_uploaded_file("resume.txt", b"some content")
    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_text(mock_file)


def test_extract_text_old_doc_format():
    """Uploading a .doc file should raise a ValueError with helpful message."""
    mock_file = make_uploaded_file("resume.doc", b"binary content")
    with pytest.raises(ValueError, match=".doc format is not supported"):
        extract_text(mock_file)


def test_extract_text_from_docx():
    """
    Tests DOCX extraction using a real in-memory DOCX file.
    We create a minimal DOCX with python-docx, write it to bytes,
    then pass those bytes to our extractor.
    """
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_paragraph("Basil Ahmed")
    doc.add_paragraph("Software Engineer with 5 years experience.")

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    result = extract_text_from_docx(docx_bytes)
    assert "Basil Ahmed" in result
    assert "Software Engineer" in result


def test_extract_text_pdf(tmp_path):
    """
    Tests PDF extraction using a real minimal PDF created with PyMuPDF.
    tmp_path is a pytest built-in fixture that gives us a temp folder.
    """
    import fitz
    pdf_path = tmp_path / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 100), "Basil Ahmed - Data Engineer")
    doc.save(str(pdf_path))
    doc.close()

    pdf_bytes = pdf_path.read_bytes()
    result = extract_text_from_pdf(pdf_bytes)
    assert "Basil Ahmed" in result
    assert "Data Engineer" in result